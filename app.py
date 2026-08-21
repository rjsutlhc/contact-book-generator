import streamlit as st
import datetime
import re
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Cm
from weasyprint import HTML

# --- 文字處理優化 ---
def apply_font_style(run, size=18):
    """將字型強制綁定為標楷體，並鎖定字級"""
    run.font.name = 'DFKai-SB'
    run.font.size = Pt(size)
    rPr = run._r.get_or_add_rPr()
    
    # 強制設定各種屬性的字型，避免 Word 自動跳回系統預設字型
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'DFKai-SB')
    rFonts.set(qn('w:hAnsi'), 'DFKai-SB')
    rFonts.set(qn('w:eastAsia'), 'DFKai-SB')
    rPr.append(rFonts)
    return run

def to_full_width(text):
    """將數字轉換為全形，以配合標楷體風格"""
    return text.translate(str.maketrans('0123456789', '０１２３４５６７８９'))

def add_run_with_tate_chu_yoko(paragraph, text):
    """產生兩位數橫向文字，並強制套用標楷體"""
    run = paragraph.add_run(text)
    # 套用樣式與字型
    apply_font_style(run, size=18)
    
    # 執行合併文字設定
    rPr = run._r.get_or_add_rPr()
    eal = OxmlElement('w:eastAsianLayout')
    eal.set(qn('w:id'), '1')
    eal.set(qn('w:combine'), '1')
    rPr.append(eal)

def add_formatted_text(paragraph, text):
    """根據數字位數，決定全形或橫向合併"""
    parts = re.split(r'(\d+)', text)
    for part in parts:
        if part.isdigit():
            if len(part) == 1:
                run = paragraph.add_run(to_full_width(part))
                apply_font_style(run)
            elif len(part) >= 2:
                # 兩位數以上執行橫向合併
                add_run_with_tate_chu_yoko(paragraph, part)
        else:
            if part:
                run = paragraph.add_run(part)
                apply_font_style(run)

# --- 檔案生成邏輯 ---
def generate_docx(date_str, items, parent_note, filename):
    doc = Document()
    
    # A4 頁面設定
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(1.91)
    section.right_margin = Cm(1.91)

    table = doc.add_table(rows=4, cols=1)
    table.style = 'Table Grid'

    for i in range(4):
        cell = table.cell(i, 0)
        # 固定高度
        tr = cell._element.getparent()
        trHeight = OxmlElement('w:trHeight')
        trHeight.set(qn('w:val'), '3500')
        trHeight.set(qn('w:hRule'), 'exact')
        tr.append(trHeight)
        
        # 強制直書
        tcPr = cell._tc.get_or_add_tcPr()
        textDir = OxmlElement('w:textDirection')
        textDir.set(qn('w:val'), 'tbRl')
        tcPr.append(textDir)
        
        p = cell.paragraphs[0]
        p.paragraph_format.line_spacing = 1.5
        p.clear()
        
        # 寫入文字
        content = [date_str, ""] + [f"{i+1} □ {item}" for i, item in enumerate(items)] + ["", "● 家長幫幫忙：", parent_note]
        for line in content:
            add_formatted_text(p, line)
            p.add_run('\n')

    doc.save(filename)

# --- Streamlit 介面 (省略重複部分，與前次相同) ---
# ... (請維持您之前的介面程式碼) ...